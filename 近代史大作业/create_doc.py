# -*- coding: utf-8 -*-
"""创建长征数字地图作品说明文档"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

def set_run_font(run, font_name='宋体', size=12, bold=False, color=None):
    """设置run的字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading1(doc, text):
    """添加一级标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(24)
    pf.space_after = Pt(12)
    run = p.add_run(text)
    set_run_font(run, '黑体', 16, bold=True, color=(192, 57, 43))
    # 下划线效果用边框实现
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), 'C0392B')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_heading2(doc, text):
    """添加二级标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, '黑体', 14, bold=True, color=(139, 0, 0))
    return p

def add_body(doc, text, indent=True, bold=False, color=None):
    """添加正文段落"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if indent:
        pf.first_line_indent = Cm(0.74)  # 两个字符的缩进
    run = p.add_run(text)
    set_run_font(run, '宋体', 12, bold=bold, color=color)
    return p

def add_bullet(doc, text):
    """添加项目符号段落"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.left_indent = Cm(0.74)
    pf.first_line_indent = Cm(-0.37)
    run = p.add_run('● ' + text)
    set_run_font(run, '宋体', 12, color=(50, 50, 50))
    return p

def add_empty(doc):
    """添加空行"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    return p

def set_cell_style(cell, bg_color=None, bold=False, align_center=False, font_size=10):
    """设置单元格样式"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # 边框
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), 'DDDDDD')
        tcBorders.append(border)
    tcPr.append(tcBorders)
    # 内边距
    tcMarins = OxmlElement('w:tcMar')
    for side, val in [('top', '80'), ('bottom', '80'), ('left', '120'), ('right', '120')]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), val)
        m.set(qn('w:type'), 'dxa')
        tcMarins.append(m)
    tcPr.append(tcMarins)
    if bg_color:
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), bg_color)
        tcPr.append(shd)
    # 垂直居中
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)
    # 文字
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
        for run in para.runs:
            run.font.size = Pt(font_size)
            run.font.bold = bold

def hex_to_rgb(h):
    h = h.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

def create_doc():
    doc = Document()
    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.header_distance = Cm(1.27)
    section.footer_distance = Cm(1.27)

    # 页眉
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run('《中国近现代史纲要》实践作业 · 长征数字地图')
    set_run_font(hr, '宋体', 9, color=(150, 150, 150))

    # 页脚页码
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run('第 ')
    set_run_font(fr, '宋体', 9, color=(150, 150, 150))
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run_page = fp.add_run()
    run_page._element.append(fldChar1)
    run_page._element.append(instrText)
    run_page._element.append(fldChar2)
    set_run_font(run_page, '宋体', 9, color=(150, 150, 150))
    fr2 = fp.add_run(' 页')
    set_run_font(fr2, '宋体', 9, color=(150, 150, 150))

    # ========== 封面 ==========
    for _ in range(4):
        add_empty(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('纪念长征胜利九十周年')
    set_run_font(run, '宋体', 14, bold=True, color=(192, 57, 43))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(12)
    run = p.add_run('长征数字地图')
    set_run_font(run, '黑体', 36, bold=True, color=(192, 57, 43))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after = Pt(8)
    run = p.add_run('——二万五千里  铁骨铸忠魂')
    set_run_font(run, '宋体', 16, color=(100, 100, 100))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(40)
    run = p.add_run('★  ★  ★  ★  ★')
    set_run_font(run, '宋体', 18, color=(192, 57, 43))

    for _ in range(6):
        add_empty(doc)

    info_items = [
        '课程名称：中国近现代史纲要',
        '作品类型：长征数字地图（HTML交互网页）',
        '学期：2025—2026学年第二学期',
        '提交时间：2026年5月',
    ]
    for item in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        run = p.add_run(item)
        set_run_font(run, '宋体', 12, color=(80, 80, 80))

    # 分页
    doc.add_page_break()

    # ========== 一、作品概述 ==========
    add_heading1(doc, '一、作品概述')
    add_body(doc, '本作品以"纪念长征胜利九十周年"为主题，以"长征数字地图"为核心概念，制作了一款基于HTML5、CSS3、SVG与JavaScript技术的数字交互式网页作品。作品通过可交互的路线地图、历史时间线、感人故事、精神传承与诗词歌赋五大板块，全方位、多维度地呈现中国工农红军二万五千里长征的壮阔历程与伟大精神。')
    add_empty(doc)
    add_body(doc, '1934年10月至1936年10月，中国工农红军三大主力（红一、红二、红四方面军）先后进行了举世闻名的二万五千里长征，历时两年，纵横11个省，行程约25000里，翻越雪山草地，突破层层封锁，最终三军胜利会师，创造了人类军事史上的壮举与奇迹。')
    add_empty(doc)
    add_body(doc, '今年是长征胜利九十周年，本作品以数字化、可视化的手段，让这段历史在新时代"活起来"，让长征精神在青年一代中薪火相传。')

    # ========== 二、作品内容结构 ==========
    add_heading1(doc, '二、作品内容结构')
    add_body(doc, '本作品共分为六大核心板块，内容丰富、结构完整：')
    add_empty(doc)

    add_heading2(doc, '（一）英雄引言板块（Hero Section）')
    add_body(doc, '以视觉冲击力强的全屏设计作为开篇，通过星空背景、动态粒子效果和红金配色，营造出庄严肃穆而又激昂向上的历史氛围。展示长征的核心数据：行程25000余里、历时368天、途经14个省份、大小战役600余次，以震撼人心的数字向读者传递长征之伟大。')

    add_heading2(doc, '（二）红军长征路线全图（Interactive Map）')
    add_body(doc, '采用SVG矢量绘图技术，手工绘制中国地图轮廓与三支红军队伍的长征路线，标注了13个重要历史节点（瑞金起点、湘江战役、遵义会议、四渡赤水、扎西会议、彝海结盟、巧渡金沙江、飞夺泸定桥、翻越雪山、穿越草地、夺取腊子口、吴起镇会师、会宁大会师）。')
    add_body(doc, '每个节点支持鼠标悬停交互，弹出包含历史日期与详细事件描述的信息框，实现了"看地图、知历史、悟精神"的沉浸式学习体验。')

    add_heading2(doc, '（三）历史事件时间线（Timeline）')
    add_body(doc, '以左右交替布局的时间线形式，详细呈现从长征开始（1934年10月）到三大主力会师（1936年10月）的10个重大历史节点，每个事件包含精确日期、事件名称与深度解读，让读者清晰把握长征的历史脉络。')

    add_heading2(doc, '（四）感人故事板块（Stories）')
    add_body(doc, '精选6个最具代表性、最广为人知的长征故事，以图文并茂的卡片形式展示：')
    add_bullet(doc, '半条被子——军民鱼水情深的最美诠释')
    add_bullet(doc, '彝海结盟——民族团结的光辉典范')
    add_bullet(doc, '七根火柴——生命最后时刻的信仰坚守')
    add_bullet(doc, '金色鱼钩——无私奉献、舍己为人的崇高精神')
    add_bullet(doc, '陈树湘断肠明志——视死如归的革命英雄主义')
    add_bullet(doc, '丰碑（军需处长）——甘当人梯、默默奉献的精神')

    add_heading2(doc, '（五）长征精神传承板块（Spirit）')
    add_body(doc, '从四个维度系统阐述伟大的长征精神：坚定革命理想信念、艰苦奋斗精神、紧密依靠群众、顾全大局无私奉献。引用毛泽东《论反对日本帝国主义的策略》（1935年）和习近平总书记纪念红军长征胜利八十周年大会讲话（2016年），以权威话语升华主题。')

    add_heading2(doc, '（六）毛泽东长征诗词板块（Poetry）')
    add_body(doc, '精选毛泽东三首长征相关词作，以精美卡片形式展示全文并附创作背景与赏析：')
    add_bullet(doc, '《七律·长征》（1935年10月）——豪情万丈，气吞山河')
    add_bullet(doc, '《忆秦娥·娄山关》（1935年2月）——雄关漫道，壮志凌云')
    add_bullet(doc, '《清平乐·六盘山》（1935年10月）——今日长缨，何日缚龙')

    # ========== 三、技术实现说明 ==========
    add_heading1(doc, '三、技术实现说明')
    add_body(doc, '本作品综合运用多种现代网页技术与AI辅助手段进行创作：')
    add_empty(doc)
    add_bullet(doc, 'HTML5：语义化结构，分区清晰，便于搜索引擎索引与无障碍阅读')
    add_bullet(doc, 'CSS3：CSS变量统一管理主题色彩，Flexbox/Grid布局，响应式设计，支持手机与PC端访问')
    add_bullet(doc, 'CSS动画：@keyframes实现节点脉冲波纹效果、页面滚动淡入效果、英雄视差滚动效果')
    add_bullet(doc, 'SVG矢量图形：手工绘制长征路线地图，使用渐变、滤镜（glow效果）、动画元素，兼容所有主流浏览器')
    add_bullet(doc, 'JavaScript：Intersection Observer API实现按需触发动画，数字滚动计数器效果，事件监听实现地图交互')
    add_bullet(doc, 'AI辅助创作：借助AI技术进行历史内容整理、文字润色、代码优化，提升作品质量与完整度')
    add_empty(doc)
    add_body(doc, '作品运行方式：直接在浏览器中打开"长征数字地图.html"文件即可访问，无需服务器环境，不依赖外部网络，完全本地运行。')

    # ========== 四、思想性说明 ==========
    add_heading1(doc, '四、思想性与价值导向说明')
    add_body(doc, '本作品始终坚持正确的政治立场，以弘扬长征精神、传承红色基因为根本出发点，具体体现在以下几个方面：')
    add_empty(doc)
    add_bullet(doc, '历史描述客观真实：所有历史事件、人物故事均严格依据权威史料，不夸大、不失实')
    add_bullet(doc, '精神传承深入人心：通过故事、引言、诗词的多元呈现方式，让长征精神从历史走进当下')
    add_bullet(doc, '体现党的领导作用：重点呈现遵义会议的历史意义，展示中国共产党从挫折中学习、在实践中成长的历程')
    add_bullet(doc, '弘扬民族精神：以彝海结盟等故事，呈现各民族团结一心的历史传统')
    add_bullet(doc, '引用习近平新时代思想：引用总书记讲话，将长征精神与新时代紧密相连')
    add_empty(doc)
    add_body(doc, '长征精神是中华民族百折不挠、自强不息的民族精神的最高表现，是中国共产党人革命精神的重要组成部分。铭记历史、传承精神、砥砺前行，是当代青年的历史责任。')

    # ========== 五、作品评分自评表 ==========
    add_heading1(doc, '五、作品评分自评表')
    add_empty(doc)

    # 创建表格
    table_data = [
        ('评分项目', '满分', '本作品对应内容', '预估得分'),
        ('学术规范（内容严谨科学无争议）', '10分', '所有史实均来源于权威历史资料，遵义会议、湘江战役等重要事件描述准确，无争议内容', '10分'),
        ('文档资料（完整提交，符合格式要求）', '10分', '提交HTML交互地图作品 + Word说明文档，命名规范，格式完整', '10分'),
        ('思想性（弘扬长征精神，传承文化）', '10分', '专设"长征精神"板块，引用毛泽东讲话、习近平总书记讲话，全面弘扬革命精神', '10分'),
        ('创新性（构思、实施、艺术表现有创新）', '30分', '数字交互地图+SVG可视化路线+节点hover事件+CSS动画+数字滚动效果，技术与艺术融合', '27分'),
        ('丰富性（内容丰富程度）', '20分', '涵盖路线地图、时间线（10个重大事件）、感人故事（6个）、精神传承、诗词（3首）五大板块', '19分'),
        ('表达风格（主题诠释合理，具备美感）', '10分', '红色主题美学设计，星空背景、金红配色，诗意与史实并重，艺术表现形式统一协调', '9分'),
        ('技术及完整度（运用AI等技术）', '10分', 'AI辅助内容创作与代码生成，运用HTML5+CSS3+SVG+JavaScript实现数字可视化', '10分'),
        ('合计', '100分', '', '95分（预估）'),
    ]

    col_widths = [3.2, 1.3, 8.0, 2.0]  # cm
    table = doc.add_table(rows=len(table_data), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(table_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.width = Cm(col_widths[j])
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in [1, 3] else WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(cell_text)
            is_header = (i == 0)
            is_total = (i == len(table_data) - 1)
            set_run_font(run, '宋体', 10,
                        bold=(is_header or is_total),
                        color=(255, 255, 255) if is_header else
                              (139, 0, 0) if is_total else (51, 51, 51))
            set_cell_style(cell,
                          bg_color='C0392B' if is_header else
                                  'F4D03F' if is_total else
                                  'FDF5F5' if j == 0 else 'FFFFFF',
                          bold=(is_header or is_total),
                          align_center=(j in [1, 3]),
                          font_size=10)

    add_empty(doc)

    # ========== 六、主要参考资料 ==========
    add_heading1(doc, '六、主要参考资料')
    add_bullet(doc, '《中国共产党历史》第一卷，中共中央党史研究室著，中共党史出版社，2011年')
    add_bullet(doc, '《红军长征史》，军事科学院军事历史研究部著，军事科学出版社，2006年')
    add_bullet(doc, '毛泽东《论反对日本帝国主义的策略》（1935年12月27日）')
    add_bullet(doc, '习近平《在纪念红军长征胜利八十周年大会上的讲话》（2016年10月21日）')
    add_bullet(doc, '人民网党史频道：http://dangshi.people.com.cn/（长征专题）')
    add_bullet(doc, '中国共产党新闻网：http://cpc.people.com.cn/（长征历史资料）')
    add_bullet(doc, '中国军网长征九十周年专题（2024年）')
    add_empty(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf = p.paragraph_format
    pf.space_before = Pt(24)
    run = p.add_run('2026年5月')
    set_run_font(run, '宋体', 12, color=(120, 120, 120))

    # 保存
    output_path = 'E:/近代史大作业/长征数字地图-作品说明文档.docx'
    doc.save(output_path)
    print(f'文档已保存到: {output_path}')

if __name__ == '__main__':
    create_doc()
